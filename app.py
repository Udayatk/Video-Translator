# Import necessary libraries
import streamlit as st
import ffmpeg
import speech_recognition as sr
from googletrans import Translator
from gtts import gTTS
from docx import Document
from docx.shared import Pt
import yt_dlp
import os
import zipfile
from io import BytesIO

# Function to extract audio from a video file
def extract_audio(video_path):
    audio_path = "extracted_audio.wav"
    ffmpeg.input(video_path).output(audio_path, acodec="pcm_s16le").run(overwrite_output=True)
    return audio_path

# Function to transcribe audio to text using Google's speech recognition
def transcribe_audio(audio_path):
    recognizer = sr.Recognizer()
    with sr.AudioFile(audio_path) as source:
        audio = recognizer.record(source)
    try:
        return recognizer.recognize_google(audio)
    except Exception:
        return ""

# Function to translate text into a target language
def translate_text(text, target_lang, translator, chunk_size=200):
    sentences = text.split(". ")
    translated_sentences = []
    for i in range(0, len(sentences), chunk_size):
        chunk = ". ".join(sentences[i:i+chunk_size])
        translated_chunk = translator.translate(chunk, dest=target_lang).text
        translated_sentences.append(translated_chunk)
    return " ".join(translated_sentences)

# Function to generate voiceover audio from translated text
def generate_voiceover(translated_text, language):
    tts = gTTS(text=translated_text, lang=language)
    voiceover_path = f"voiceover_{language}.mp3"
    tts.save(voiceover_path)
    return voiceover_path

# Function to replace original video audio with translated voiceover
def replace_audio_with_translation(video_path, voiceover_path, lang):
    output_path = f"final_video_{lang}.mp4"
    try:
        video = ffmpeg.input(video_path).video
        voiceover = ffmpeg.input(voiceover_path).audio
        ffmpeg.output(video, voiceover, output_path, vcodec="libx264", acodec="aac", strict='experimental').run(overwrite_output=True)
    except Exception:
        return None
    return output_path

# Function to download a video from YouTube using yt_dlp
def download_youtube_video(url):
    output_file = "downloaded_video.mp4"
    if os.path.exists(output_file):
        os.remove(output_file)
    ydl_opts = {
        'format': 'bestvideo+bestaudio/best',
        'outtmpl': output_file,
        'merge_output_format': 'mp4',
        'noplaylist': True
    }
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        ydl.download([url])
    return output_file

# Function to create a Word document with original and translated texts
def create_translation_document(original_text, translations, languages):
    doc = Document()
    doc.add_heading('Translations', level=1)
    doc.add_heading('Original Text:', level=2)
    paragraph = doc.add_paragraph(original_text)
    paragraph.runs[0].font.name = "Nirmala UI"
    paragraph.runs[0].font.size = Pt(12)
    for lang, translation in zip(languages, translations):
        doc.add_heading(f'Translation in {lang}:', level=2)
        para = doc.add_paragraph(translation)
        para.runs[0].font.name = "Nirmala UI"
        para.runs[0].font.size = Pt(12)
    doc_path = "translations.docx"
    doc.save(doc_path)
    return doc_path

# Function to get the duration of a video file
def get_video_duration(video_path):
    probe = ffmpeg.probe(video_path)
    return float(probe['format']['duration'])

# Streamlit title
st.title("🎙️ Translate & Download Your Video")

# File uploader for video
video_file = st.file_uploader("📂 Upload a video file", type=["mp4", "mov"])
# Text input for YouTube URL
video_url = st.text_input("🔗 Or paste a YouTube URL:")

# Multi-select for choosing translation languages
target_languages = st.multiselect(
    "🌐 Select translation languages",
    ["hi", "te", "ta", "kn", "mr", "ur", "ml", "pa", "gu"]
)

# Dictionary to map language codes to names
language_names = {
    "hi": "Hindi", "te": "Telugu", "ta": "Tamil", "kn": "Kannada",
    "mr": "Marathi", "ur": "Urdu", "ml": "Malayalam", "pa": "Punjabi", "gu": "Gujarati"
}

# Button to trigger the translation and generation process
if st.button("Translate & Generate"):
    if not (video_file or video_url):
        st.warning("⚠️ Please upload a video or provide a URL.")
        st.stop()

    with st.spinner("Processing..."):
        # Save uploaded video to disk or download from YouTube
        if video_file:
            video_path = "uploaded_video.mp4"
            with open(video_path, "wb") as f:
                f.write(video_file.getbuffer())
        else:
            video_path = download_youtube_video(video_url)

        # Initialize translator
        translator = Translator()
        # Extract audio and transcribe it
        audio_path = extract_audio(video_path)
        transcript = transcribe_audio(audio_path)
        translations = []

        translated_video_paths = []
        # Loop through selected languages for translation and voiceover
        for lang_code in target_languages:
            translated_text = translate_text(transcript, lang_code, translator)
            translations.append(translated_text)

            voiceover_path = generate_voiceover(translated_text, lang_code)
            translated_video = replace_audio_with_translation(video_path, voiceover_path, lang_code)
            if translated_video:
                translated_video_paths.append((lang_code, translated_video))

        # Create Word document with translations
        doc_path = create_translation_document(transcript, translations, target_languages)

        # Bundle videos and doc into a zip file
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zipf:
            for lang_code, video_path in translated_video_paths:
                zipf.write(video_path, arcname=os.path.basename(video_path))
            zipf.write(doc_path, arcname="translations.docx")

        zip_buffer.seek(0)

    # Show success message
    st.success("✅ Translation completed and files are ready!")

    # Download button for the ZIP file
    st.download_button(
        label="📦 Download All Files (ZIP)",
        data=zip_buffer,
        file_name="translated_outputs.zip",
        mime="application/zip"
    )
